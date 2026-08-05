"""
Chargeur de fichiers DOCX pour le moteur d'intelligence documentaire GalSen IA.
"""

import os
from .document_loader_factory import BaseDocumentLoader
from .types import DocumentItem, DocumentType
import logging

logger = logging.getLogger(__name__)


class DOCXLoader(BaseDocumentLoader):
    """Chargeur pour les fichiers DOCX (.docx)."""

    # Extensions supportées
    supported_extensions = ['.docx']

    def load(self, source: str, **kwargs) -> DocumentItem:
        """
        Charge un fichier DOCX et retourne un DocumentItem.

        Args:
            source: Chemin vers le fichier DOCX
            **kwargs: Options supplémentaires

        Returns:
            DocumentItem représentant le fichier chargé
        """
        if not os.path.isfile(source):
            raise FileNotFoundError(f"Fichier non trouvé: {source}")

        try:
            import docx
        except ImportError:
            message = (
                "python-docx est requis pour charger les fichiers DOCX. "
                "Installez-le avec: pip install python-docx"
            )
            logger.error(message)
            raise ImportError(message)

        # Charger le document DOCX
        doc = docx.Document(source)

        # Extraire le texte paragraphe par paragraphe
        text_content = []
        for paragraph in doc.paragraphs:
            text_content.append(paragraph.text)
        # Aussi extraire le texte des tableaux
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text_content.append(cell.text)

        full_text = '\n'.join(text_content)

        # Extraire les métadonnées du document
        core_props = doc.core_properties
        metadata = {
            "size": os.path.getsize(source),
            "created": os.path.getctime(source),
            "modified": os.path.getmtime(source),
            "source": os.path.abspath(source),
            "filename": os.path.basename(source),
            "extension": ".docx",
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables)
        }

        if core_props.author:
            metadata["author"] = core_props.author
        if core_props.created:
            metadata["created"] = core_props.created.timestamp()
        if core_props.modified:
            metadata["modified"] = core_props.modified.timestamp()
        if core_props.title:
            metadata["title"] = core_props.title
        if core_props.subject:
            metadata["subject"] = core_props.subject
        if hasattr(core_props, 'keywords') and core_props.keywords:
            metadata["keywords"] = core_props.keywords

        # Utiliser le titre du document ou le nom du fichier
        title = metadata.get("title") or os.path.splitext(os.path.basename(source))[0]

        # Créer l'élément de document
        document = DocumentItem(
            document_id="",
            document_type=DocumentType.DOCX,
            title=title,
            content=full_text,
            metadata=metadata
        )

        return document

    def supports_type(self, document_type: str) -> bool:
        """Vérifie si le chargeur supporte le type de document donné."""
        return document_type in [
            DocumentType.DOCX.value,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        ]